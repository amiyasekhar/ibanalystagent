"""
Generate Word documents (DOCX) from structured memo content.
Supports investment memos, reports, and formatted documents.
"""

import sys
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def create_memo_docx(memo_data: dict, output_path: str) -> dict:
    """
    Generate a Word document from memo data.

    Args:
        memo_data: Dictionary with structure:
            {
                "title": str,
                "date": str (optional),
                "author": str (optional),
                "sections": [
                    {
                        "heading": str,
                        "level": int (1-3),
                        "content": str or list of paragraphs
                    }
                ]
            }
        output_path: Path to save the generated DOCX file

    Returns:
        {"ok": True, "path": str} or {"ok": False, "error": str}
    """
    try:
        doc = Document()

        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        title = doc.add_heading(memo_data.get("title", "Investment Memo"), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata (date, author)
        if memo_data.get("date"):
            date_para = doc.add_paragraph(f"Date: {memo_data['date']}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_para.runs[0].font.size = Pt(11)

        if memo_data.get("author"):
            author_para = doc.add_paragraph(f"Prepared by: {memo_data['author']}")
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_para.runs[0].font.size = Pt(11)

        # Add spacing
        doc.add_paragraph()

        # Add sections
        sections_list = memo_data.get("sections", [])
        for section in sections_list:
            heading = section.get("heading", "")
            level = section.get("level", 1)
            content = section.get("content", "")

            # Add heading
            if heading:
                doc.add_heading(heading, level=level)

            # Add content
            if isinstance(content, list):
                for para_text in content:
                    para = doc.add_paragraph(para_text)
                    para.style = 'Normal'
                    para.paragraph_format.space_after = Pt(12)
            else:
                para = doc.add_paragraph(content)
                para.style = 'Normal'
                para.paragraph_format.space_after = Pt(12)

        # Save document
        doc.save(output_path)

        return {
            "ok": True,
            "path": output_path
        }

    except Exception as e:
        return {
            "ok": False,
            "error": str(e)
        }


def extract_text_from_docx(docx_path: str) -> dict:
    """
    Extract text content from a Word document.

    Args:
        docx_path: Path to DOCX file

    Returns:
        {"ok": True, "text": str, "sections": list} or {"ok": False, "error": str}
    """
    try:
        doc = Document(docx_path)

        full_text = []
        sections = []
        current_section = None

        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                para = element
                para_text = ''.join(node.text for node in para.iter() if hasattr(node, 'text'))

                # Check if it's a heading
                if para.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle'):
                    style = para.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                    if 'Heading' in style:
                        # Save previous section
                        if current_section:
                            sections.append(current_section)

                        # Start new section
                        current_section = {
                            "heading": para_text.strip(),
                            "content": []
                        }
                        full_text.append(f"\n## {para_text.strip()}\n")
                        continue

                # Regular paragraph
                if para_text.strip():
                    full_text.append(para_text.strip())
                    if current_section:
                        current_section["content"].append(para_text.strip())

        # Add last section
        if current_section:
            sections.append(current_section)

        return {
            "ok": True,
            "text": "\n".join(full_text),
            "sections": sections
        }

    except Exception as e:
        return {
            "ok": False,
            "error": str(e)
        }


def main():
    """CLI interface for document export."""
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "No command provided"}))
        sys.exit(1)

    command = sys.argv[1]

    try:
        if command == "create":
            # Read memo data from stdin
            memo_data = json.loads(sys.stdin.read())
            output_path = memo_data.get("outputPath", "output.docx")

            result = create_memo_docx(memo_data, output_path)
            print(json.dumps(result))

        elif command == "extract":
            # Extract text from existing DOCX
            if len(sys.argv) < 3:
                print(json.dumps({"ok": False, "error": "No file path provided"}))
                sys.exit(1)

            docx_path = sys.argv[2]
            result = extract_text_from_docx(docx_path)
            print(json.dumps(result))

        else:
            print(json.dumps({"ok": False, "error": f"Unknown command: {command}"}))
            sys.exit(1)

    except Exception as e:
        print(json.dumps({"ok": False, "error": str(e)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
