"""
Extract text and structure from precedent investment memos (PDF or DOCX).
Analyzes document structure: sections, headings, formatting patterns.
"""

import sys
import json
import fitz  # PyMuPDF
from docx import Document
from pathlib import Path


def extract_from_pdf(pdf_path: str) -> dict:
    """
    Extract text and structure from a PDF precedent memo.

    Args:
        pdf_path: Path to PDF file

    Returns:
        {"ok": True, "text": str, "sections": list, "metadata": dict}
    """
    try:
        doc = fitz.open(pdf_path)
        full_text = []
        sections = []
        current_section = None

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()

            # Simple section detection: lines in ALL CAPS or title case with short length
            lines = text.split('\n')
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    continue

                # Detect potential headings (heuristic)
                is_heading = (
                    len(line_stripped) < 80 and
                    (line_stripped.isupper() or
                     (line_stripped[0].isupper() and ':' in line_stripped) or
                     line_stripped.endswith(':'))
                )

                if is_heading and len(line_stripped) > 5:
                    # Save previous section
                    if current_section and current_section.get("content"):
                        sections.append(current_section)

                    # Start new section
                    current_section = {
                        "heading": line_stripped.rstrip(':'),
                        "content": [],
                        "page": page_num + 1
                    }
                    full_text.append(f"\n## {line_stripped}\n")
                else:
                    # Regular content
                    full_text.append(line_stripped)
                    if current_section:
                        current_section["content"].append(line_stripped)

        # Add last section
        if current_section and current_section.get("content"):
            sections.append(current_section)

        doc.close()

        return {
            "ok": True,
            "text": "\n".join(full_text),
            "sections": sections,
            "metadata": {
                "total_pages": len(doc),
                "filename": Path(pdf_path).name
            }
        }

    except Exception as e:
        return {
            "ok": False,
            "error": f"PDF extraction failed: {str(e)}"
        }


def extract_from_docx(docx_path: str) -> dict:
    """
    Extract text and structure from a DOCX precedent memo.

    Args:
        docx_path: Path to DOCX file

    Returns:
        {"ok": True, "text": str, "sections": list, "metadata": dict}
    """
    try:
        doc = Document(docx_path)

        full_text = []
        sections = []
        current_section = None

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            # Check if it's a heading
            if para.style.name.startswith('Heading'):
                # Save previous section
                if current_section and current_section.get("content"):
                    sections.append(current_section)

                # Start new section
                level = int(para.style.name.replace('Heading ', '')) if len(para.style.name) > 7 else 1
                current_section = {
                    "heading": para_text,
                    "level": level,
                    "content": []
                }
                full_text.append(f"\n{'#' * level} {para_text}\n")
            else:
                # Regular paragraph
                full_text.append(para_text)
                if current_section:
                    current_section["content"].append(para_text)

        # Add last section
        if current_section and current_section.get("content"):
            sections.append(current_section)

        return {
            "ok": True,
            "text": "\n".join(full_text),
            "sections": sections,
            "metadata": {
                "filename": Path(docx_path).name,
                "total_paragraphs": len(doc.paragraphs)
            }
        }

    except Exception as e:
        return {
            "ok": False,
            "error": f"DOCX extraction failed: {str(e)}"
        }


def extract_precedent(file_path: str) -> dict:
    """
    Extract text from precedent memo (auto-detect PDF or DOCX).

    Args:
        file_path: Path to precedent file

    Returns:
        Extracted content with structure
    """
    file_path_obj = Path(file_path)

    if not file_path_obj.exists():
        return {
            "ok": False,
            "error": f"File not found: {file_path}"
        }

    extension = file_path_obj.suffix.lower()

    if extension == '.pdf':
        return extract_from_pdf(file_path)
    elif extension in ['.docx', '.doc']:
        return extract_from_docx(file_path)
    else:
        return {
            "ok": False,
            "error": f"Unsupported file type: {extension}. Use PDF or DOCX."
        }


def main():
    """CLI interface for precedent extraction."""
    if len(sys.argv) < 2:
        print(json.dumps({"ok": False, "error": "No file path provided"}))
        sys.exit(1)

    file_path = sys.argv[1]
    result = extract_precedent(file_path)
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
